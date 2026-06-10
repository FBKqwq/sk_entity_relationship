from __future__ import annotations

import argparse
import json
import math
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


LABEL_CN = {
    "sub_diseases": "疾病/亚型",
    "symptoms": "症状/临床表现",
    "tests": "检查/检验",
    "treatments": "治疗措施",
    "plans": "治疗方案",
    "etiologies": "病因",
    "pathogeneses": "发病机制",
}


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if not path.exists() or path.stat().st_size == 0:
        return rows
    with path.open("r", encoding="utf-8-sig") as f:
        for line_no, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError as exc:
                raise ValueError(f"{path} line {line_no} is not valid JSON") from exc
    return rows


def strip_suffix(name: str) -> str:
    for suffix in (".chunk_label_result.jsonl", ".lv1_lf_outputs.jsonl"):
        if name.endswith(suffix):
            return name[: -len(suffix)]
    return Path(name).stem


def try_official_lfanalysis_available() -> tuple[bool, str]:
    try:
        import snorkel  # type: ignore
        from snorkel.labeling import LFAnalysis  # noqa: F401

        return True, getattr(snorkel, "__version__", "unknown")
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)


def safe_pct(numerator: float, denominator: float) -> float:
    return 0.0 if denominator == 0 else numerator / denominator


def add_doc_name(rows: list[dict[str, Any]], doc_name: str) -> None:
    for row in rows:
        row["doc_name"] = doc_name
        row["row_key"] = f"{doc_name}::{row.get('chunk_id', '')}::{row.get('label', '')}"


def load_dataset(input_dir: Path) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    files = list(input_dir.glob("*.jsonl"))
    by_doc: dict[str, dict[str, Path]] = defaultdict(dict)
    for path in files:
        doc_name = strip_suffix(path.name)
        if path.name.endswith(".chunk_label_result.jsonl"):
            by_doc[doc_name]["chunk"] = path
        elif path.name.endswith(".lv1_lf_outputs.jsonl"):
            by_doc[doc_name]["lf"] = path

    chunk_rows: list[dict[str, Any]] = []
    lf_rows: list[dict[str, Any]] = []
    doc_rows: list[dict[str, Any]] = []
    for doc_name in sorted(by_doc):
        chunk_path = by_doc[doc_name].get("chunk")
        lf_path = by_doc[doc_name].get("lf")
        chunk_size = chunk_path.stat().st_size if chunk_path and chunk_path.exists() else 0
        lf_size = lf_path.stat().st_size if lf_path and lf_path.exists() else 0
        chunks = read_jsonl(chunk_path) if chunk_path else []
        lfs = read_jsonl(lf_path) if lf_path else []
        add_doc_name(chunks, doc_name)
        add_doc_name(lfs, doc_name)
        chunk_rows.extend(chunks)
        lf_rows.extend(lfs)
        doc_rows.append(
            {
                "doc_name": doc_name,
                "chunk_file_bytes": chunk_size,
                "lf_file_bytes": lf_size,
                "chunk_label_rows": len(chunks),
                "lf_rows": len(lfs),
                "is_empty_output": int(chunk_size == 0 or lf_size == 0 or not chunks or not lfs),
            }
        )
    return pd.DataFrame(chunk_rows), pd.DataFrame(lf_rows), pd.DataFrame(doc_rows)


def build_lf_summary(lf_df: pd.DataFrame) -> pd.DataFrame:
    if lf_df.empty:
        return pd.DataFrame()

    rows_total = lf_df["row_key"].nunique()
    matrix_votes = defaultdict(dict)
    for row in lf_df.itertuples(index=False):
        matrix_votes[getattr(row, "row_key")][getattr(row, "lf_name")] = int(getattr(row, "vote"))

    lf_names = sorted(lf_df["lf_name"].dropna().unique())
    summary_rows: list[dict[str, Any]] = []
    for lf_name in lf_names:
        sub = lf_df[lf_df["lf_name"] == lf_name].copy()
        positives = int((sub["vote"] > 0).sum())
        negatives = int((sub["vote"] < 0).sum())
        abstains = int((sub["vote"] == 0).sum())
        non_abstain = positives + negatives
        covered_keys = set(sub.loc[sub["vote"] != 0, "row_key"])
        positive_keys = set(sub.loc[sub["vote"] > 0, "row_key"])
        overlap_count = 0
        conflict_count = 0
        for key in covered_keys:
            own_vote = int(sub.loc[sub["row_key"] == key, "vote"].iloc[0])
            other_votes = [
                vote
                for other_lf, vote in matrix_votes[key].items()
                if other_lf != lf_name and vote != 0
            ]
            if other_votes:
                overlap_count += 1
            if own_vote != 0 and any(vote != own_vote for vote in other_votes):
                conflict_count += 1
        polarity = []
        if negatives:
            polarity.append("negative")
        if positives:
            polarity.append("positive")
        summary_rows.append(
            {
                "LF": lf_name,
                "Polarity": ", ".join(polarity) if polarity else "abstain only",
                "Coverage": safe_pct(non_abstain, rows_total),
                "PositiveRate": safe_pct(positives, rows_total),
                "Overlap": safe_pct(overlap_count, rows_total),
                "Conflict": safe_pct(conflict_count, rows_total),
                "PositiveVotes": positives,
                "Abstains": abstains,
                "MeanConfidenceWhenPositive": float(sub.loc[sub["vote"] > 0, "confidence"].mean())
                if positives
                else 0.0,
                "MeanCountWhenPositive": float(sub.loc[sub["vote"] > 0, "count"].mean())
                if positives
                else 0.0,
                "RowsSeen": int(len(sub)),
            }
        )
    return pd.DataFrame(summary_rows).sort_values("Coverage", ascending=False)


def build_official_lfanalysis_summary(lf_df: pd.DataFrame) -> pd.DataFrame:
    """Run Snorkel's official LFAnalysis on the persisted Lv1 matrix."""

    from snorkel.labeling import LFAnalysis, LabelingFunction  # type: ignore

    lf_names = sorted(lf_df["lf_name"].dropna().unique())
    row_keys = sorted(lf_df["row_key"].dropna().unique())
    row_index = {key: idx for idx, key in enumerate(row_keys)}
    lf_index = {name: idx for idx, name in enumerate(lf_names)}
    matrix = np.full((len(row_keys), len(lf_names)), -1, dtype=int)
    for row in lf_df.itertuples(index=False):
        vote = int(getattr(row, "vote"))
        if vote > 0:
            value = 1
        elif vote < 0:
            value = 0
        else:
            value = -1
        matrix[row_index[getattr(row, "row_key")], lf_index[getattr(row, "lf_name")]] = value

    dummy_lfs = [
        LabelingFunction(name=name, f=lambda x, _name=name: -1)
        for name in lf_names
    ]
    official = LFAnalysis(L=matrix, lfs=dummy_lfs).lf_summary()
    official = official.reset_index().rename(columns={"index": "LF"})
    official.columns = [str(col) for col in official.columns]
    if "Polarity" in official.columns:
        def clean_polarity(value: Any) -> str:
            if isinstance(value, np.ndarray):
                values = value.tolist()
            elif isinstance(value, (list, tuple, set)):
                values = list(value)
            else:
                values = [value]
            names = []
            for item in values:
                raw = int(item) if isinstance(item, (np.integer, int)) else item
                if raw == 1:
                    names.append("positive")
                elif raw == 0:
                    names.append("negative")
                elif raw == -1:
                    names.append("abstain")
                else:
                    names.append(str(raw))
            return ", ".join(names)

        official["Polarity"] = official["Polarity"].map(clean_polarity)
    return official.sort_values("Coverage", ascending=False)


def build_label_summary(chunk_df: pd.DataFrame) -> pd.DataFrame:
    if chunk_df.empty:
        return pd.DataFrame()
    rows = []
    for label, sub in chunk_df.groupby("label"):
        accepted = sub[(sub["present"] == True) & (sub["status"] == "accepted")]  # noqa: E712
        weak = sub[sub["status"] == "weak"]
        rejected = sub[sub["status"] == "rejected"]
        no_support = accepted[
            accepted["supporting_lfs"].apply(lambda x: len(x) if isinstance(x, list) else 0) == 0
        ]
        no_evidence = accepted[
            accepted["evidence_texts"].apply(lambda x: len(x) if isinstance(x, list) else 0) == 0
        ]
        rows.append(
            {
                "label": label,
                "label_cn": LABEL_CN.get(label, label),
                "total_rows": int(len(sub)),
                "accepted_rows": int(len(accepted)),
                "accepted_rate": safe_pct(len(accepted), len(sub)),
                "weak_rows": int(len(weak)),
                "rejected_rows": int(len(rejected)),
                "avg_confidence": float(accepted["confidence"].mean()) if len(accepted) else 0.0,
                "avg_predicted_count": float(accepted["predicted_count"].mean()) if len(accepted) else 0.0,
                "accepted_without_support": int(len(no_support)),
                "accepted_without_evidence": int(len(no_evidence)),
            }
        )
    return pd.DataFrame(rows).sort_values("accepted_rows", ascending=False)


def build_doc_summary(chunk_df: pd.DataFrame, doc_df: pd.DataFrame) -> pd.DataFrame:
    if chunk_df.empty:
        return doc_df.copy()
    accepted = chunk_df[(chunk_df["present"] == True) & (chunk_df["status"] == "accepted")]  # noqa: E712
    chunks = chunk_df.groupby("doc_name")["chunk_id"].nunique().rename("chunk_count")
    accepted_rows = accepted.groupby("doc_name").size().rename("accepted_label_rows")
    avg_conf = accepted.groupby("doc_name")["confidence"].mean().rename("avg_accepted_confidence")
    labels_per_doc = accepted.groupby("doc_name")["label"].nunique().rename("accepted_label_types")
    out = doc_df.merge(chunks, on="doc_name", how="left")
    out = out.merge(accepted_rows, on="doc_name", how="left")
    out = out.merge(avg_conf, on="doc_name", how="left")
    out = out.merge(labels_per_doc, on="doc_name", how="left")
    for col in ["chunk_count", "accepted_label_rows", "accepted_label_types"]:
        out[col] = out[col].fillna(0).astype(int)
    out["avg_accepted_confidence"] = out["avg_accepted_confidence"].fillna(0.0)
    out["labels_per_chunk"] = out.apply(
        lambda r: safe_pct(r["accepted_label_rows"], r["chunk_count"]), axis=1
    )
    return out.sort_values("accepted_label_rows", ascending=False)


def save_table(df: pd.DataFrame, path: Path) -> None:
    df.to_csv(path, index=False, encoding="utf-8-sig")


def configure_plot() -> None:
    available = {font.name for font in font_manager.fontManager.ttflist}
    for candidate in ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Arial Unicode MS"]:
        if candidate in available:
            plt.rcParams["font.family"] = [candidate]
            break
    else:
        plt.rcParams["font.family"] = ["DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 140


def chart_label_distribution(label_summary: pd.DataFrame, out: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    data = label_summary.sort_values("accepted_rows", ascending=True)
    y = np.arange(len(data))
    colors = ["#2E74B5" if v >= data["accepted_rows"].median() else "#8FB3D9" for v in data["accepted_rows"]]
    ax.barh(y, data["accepted_rows"], color=colors)
    ax.set_yticks(y, data["label"].map(lambda x: LABEL_CN.get(x, x)))
    ax.set_xlabel("Accepted chunk-label rows")
    ax.set_title("Lv1 Accepted Label Distribution")
    for i, value in enumerate(data["accepted_rows"]):
        ax.text(value, i, f" {int(value)}", va="center", fontsize=9)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def chart_lf_coverage(lf_summary: pd.DataFrame, out: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.2, 4.3))
    data = lf_summary.sort_values("Coverage", ascending=True)
    y = np.arange(len(data))
    ax.barh(y, data["Coverage"] * 100, color="#1F4D78")
    ax.set_yticks(y, data["LF"])
    ax.set_xlabel("Coverage (%)")
    ax.set_title("LFAnalysis-style Coverage by Labeling Function")
    for i, value in enumerate(data["Coverage"] * 100):
        ax.text(value, i, f" {value:.1f}%", va="center", fontsize=9)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def chart_lf_label_heatmap(lf_df: pd.DataFrame, out: Path) -> None:
    pivot = (
        lf_df.assign(positive=(lf_df["vote"] > 0).astype(int))
        .groupby(["lf_name", "label"])["positive"]
        .mean()
        .unstack(fill_value=0.0)
    )
    pivot = pivot.reindex(index=sorted(pivot.index), columns=sorted(pivot.columns))
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    im = ax.imshow(pivot.values * 100, cmap="Blues", aspect="auto", vmin=0)
    ax.set_xticks(np.arange(len(pivot.columns)))
    ax.set_xticklabels([LABEL_CN.get(c, c) for c in pivot.columns], rotation=35, ha="right")
    ax.set_yticks(np.arange(len(pivot.index)))
    ax.set_yticklabels(pivot.index)
    ax.set_title("Positive Vote Rate by LF and Label")
    for i in range(pivot.shape[0]):
        for j in range(pivot.shape[1]):
            value = pivot.values[i, j] * 100
            ax.text(j, i, f"{value:.0f}", ha="center", va="center", fontsize=8, color="#1B1B1B")
    cbar = fig.colorbar(im, ax=ax)
    cbar.set_label("% positive votes")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def chart_doc_concentration(doc_summary: pd.DataFrame, out: Path) -> pd.DataFrame:
    non_empty = doc_summary[doc_summary["chunk_count"] > 0].copy()
    top = non_empty.nlargest(12, "accepted_label_rows").copy()
    top["doc_code"] = [f"D{i:02d}" for i in range(1, len(top) + 1)]
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    data = top.sort_values("accepted_label_rows", ascending=True)
    y = np.arange(len(data))
    ax.barh(y, data["accepted_label_rows"], color="#6A8CAF")
    ax.set_yticks(y, data["doc_code"])
    ax.set_xlabel("Accepted chunk-label rows")
    ax.set_title("Top Documents by Accepted Signals")
    for i, value in enumerate(data["accepted_label_rows"]):
        ax.text(value, i, f" {int(value)}", va="center", fontsize=8)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return top[["doc_code", "doc_name", "chunk_count", "accepted_label_rows", "labels_per_chunk"]]


def chart_quality_flags(label_summary: pd.DataFrame, out: Path) -> None:
    data = label_summary.sort_values("accepted_without_support", ascending=True)
    fig, ax = plt.subplots(figsize=(8.2, 4.4))
    y = np.arange(len(data))
    ax.barh(y, data["accepted_without_support"], label="No supporting LF", color="#C46A4A")
    ax.barh(
        y,
        data["accepted_without_evidence"],
        left=data["accepted_without_support"],
        label="No evidence text",
        color="#E4B363",
    )
    ax.set_yticks(y, data["label"].map(lambda x: LABEL_CN.get(x, x)))
    ax.set_xlabel("Accepted rows")
    ax.set_title("Accepted Rows Needing Evidence QA")
    ax.legend(loc="lower right")
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: Any, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if isinstance(text, (int, float)) else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], max_rows: int = 20) -> None:
    table_df = df.loc[:, columns].head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.2f}" if abs(value) >= 1 else f"{value:.1%}"
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_picture_with_caption(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(6.25))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def add_metric_table(doc: Document, metrics: dict[str, Any]) -> None:
    df = pd.DataFrame(
        [
            {"metric": "文档总数", "value": metrics["doc_total"]},
            {"metric": "有效文档数", "value": metrics["doc_non_empty"]},
            {"metric": "空输出/无有效输出文档", "value": metrics["doc_empty"]},
            {"metric": "chunk 数", "value": metrics["chunk_total"]},
            {"metric": "chunk-label 结果行", "value": metrics["chunk_label_rows"]},
            {"metric": "LF 输出行", "value": metrics["lf_rows"]},
            {"metric": "接受标签行", "value": metrics["accepted_rows"]},
            {"metric": "接受率", "value": f"{metrics['accepted_rate']:.1%}"},
            {"metric": "平均接受置信度", "value": f"{metrics['avg_accepted_confidence']:.3f}"},
            {"metric": "官方 Snorkel LFAnalysis", "value": metrics["official_lfanalysis"]},
        ]
    )
    add_table(doc, df, ["metric", "value"], ["指标", "结果"], max_rows=20)


def build_report(
    *,
    out_docx: Path,
    input_dir: Path,
    charts_dir: Path,
    chunk_df: pd.DataFrame,
    lf_df: pd.DataFrame,
    doc_summary: pd.DataFrame,
    label_summary: pd.DataFrame,
    lf_summary: pd.DataFrame,
    official_lf_summary: pd.DataFrame | None,
    top_docs: pd.DataFrame,
    official_available: bool,
    official_reason: str,
    chart_paths: dict[str, Path],
) -> None:
    accepted = chunk_df[(chunk_df["present"] == True) & (chunk_df["status"] == "accepted")] if not chunk_df.empty else chunk_df  # noqa: E712
    chunk_total = chunk_df[["doc_name", "chunk_id"]].drop_duplicates().shape[0] if not chunk_df.empty else 0
    metrics = {
        "doc_total": int(len(doc_summary)),
        "doc_non_empty": int((doc_summary["is_empty_output"] == 0).sum()),
        "doc_empty": int((doc_summary["is_empty_output"] == 1).sum()),
        "chunk_total": int(chunk_total),
        "chunk_label_rows": int(len(chunk_df)),
        "lf_rows": int(len(lf_df)),
        "accepted_rows": int(len(accepted)),
        "accepted_rate": safe_pct(len(accepted), len(chunk_df)),
        "avg_accepted_confidence": float(accepted["confidence"].mean()) if len(accepted) else 0.0,
        "official_lfanalysis": f"可用，版本 {official_reason}" if official_available else "未能导入，本报告使用等价指标口径",
    }

    doc = Document()
    set_doc_styles(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Lv1 Chunk 标签结果与标注函数分析报告")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    p = doc.add_paragraph(f"数据目录：{input_dir}")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    p = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("1. 结论概览", level=1)
    add_bullet(
        doc,
        f"本次共发现 {metrics['doc_total']} 份文档输出，其中 {metrics['doc_non_empty']} 份有有效 Lv1 标签结果，"
        f"{metrics['doc_empty']} 份为空输出或无有效记录，建议后续复跑或检查 PDF 解析链路。",
    )
    add_bullet(
        doc,
        f"有效结果覆盖 {metrics['chunk_total']} 个 chunk，形成 {metrics['chunk_label_rows']} 行 chunk-label 判断；"
        f"其中 {metrics['accepted_rows']} 行被接受，整体接受率为 {metrics['accepted_rate']:.1%}。",
    )
    if math.isclose(metrics["accepted_rate"], 1.0):
        add_bullet(
            doc,
            "需要重点说明：本批有效 chunk 对 7 个 Lv1 标签全部被 accepted。"
            "这说明当前结果更适合作为“高召回候选标签筛选”，尚不能直接视为精筛后的最终标签。"
        )
    add_bullet(
        doc,
        "Lv1 的主要信号来自 LLM 语义判断、证据锚定、边界数量判断和词典 LF；医学模式 LF 对症状、治疗、检查等标签有辅助价值。",
    )
    if not official_available:
        add_bullet(
            doc,
            f"当前环境未能导入官方 Snorkel 包，因此未直接运行 LFAnalysis；导入/安装失败原因：{official_reason}。"
            "报告中的 LF 覆盖率、重叠率、冲突率按 Snorkel LFAnalysis 的核心口径从落盘矩阵复算。",
        )

    doc.add_heading("2. 核心统计", level=1)
    add_metric_table(doc, metrics)

    doc.add_heading("3. 标签结果分布", level=1)
    doc.add_paragraph(
        "下图展示各类实体标签在 chunk 级别被接受的数量。数量越高，说明该类标签在专家共识文本中越常被当前 Lv1 规则和 LLM 信号共同识别。"
    )
    add_picture_with_caption(doc, chart_paths["label_distribution"], "图 1：各标签 accepted 结果数量分布")
    add_table(
        doc,
        label_summary.assign(
            accepted_rate=label_summary["accepted_rate"].map(lambda x: f"{x:.1%}"),
            avg_confidence=label_summary["avg_confidence"].map(lambda x: f"{x:.3f}"),
            avg_predicted_count=label_summary["avg_predicted_count"].map(lambda x: f"{x:.2f}"),
        ),
        [
            "label_cn",
            "accepted_rows",
            "accepted_rate",
            "avg_confidence",
            "avg_predicted_count",
            "accepted_without_support",
            "accepted_without_evidence",
        ],
        ["标签", "接受数", "接受率", "平均置信度", "平均数量", "无支持LF", "无证据"],
        max_rows=10,
    )

    doc.add_heading("4. LFAnalysis 视角", level=1)
    doc.add_paragraph(
        "LFAnalysis 关注每个标注函数是否真正贡献信号、是否与其他 LF 同时命中，以及是否产生冲突。"
        "本批 Lv1 输出以 positive/abstain 为主，几乎没有 negative vote，因此冲突率低是预期现象；更值得关注的是覆盖率和重叠率。"
    )
    add_picture_with_caption(doc, chart_paths["lf_coverage"], "图 2：各 LF 的覆盖率")
    add_picture_with_caption(doc, chart_paths["lf_heatmap"], "图 3：不同 LF 对不同标签的正投票比例")
    lf_table = lf_summary.copy()
    for col in ["Coverage", "PositiveRate", "Overlap", "Conflict"]:
        lf_table[col] = lf_table[col].map(lambda x: f"{x:.1%}")
    lf_table["MeanConfidenceWhenPositive"] = lf_table["MeanConfidenceWhenPositive"].map(lambda x: f"{x:.3f}")
    add_table(
        doc,
        lf_table,
        ["LF", "Polarity", "Coverage", "Overlap", "Conflict", "PositiveVotes", "MeanConfidenceWhenPositive"],
        ["LF", "极性", "覆盖率", "重叠率", "冲突率", "正票数", "正票均置信度"],
        max_rows=12,
    )
    if official_lf_summary is not None and not official_lf_summary.empty:
        doc.add_paragraph("下表为官方 Snorkel LFAnalysis.lf_summary() 的直接输出摘录，用于和上面的解释性统计互相校验。")
        official_show = official_lf_summary.copy()
        for col in ["Polarity"]:
            if col in official_show.columns:
                official_show[col] = official_show[col].map(lambda x: ", ".join(map(str, x)) if isinstance(x, (list, tuple, np.ndarray)) else str(x))
        for col in ["Coverage", "Overlaps", "Conflicts"]:
            if col in official_show.columns:
                official_show[col] = official_show[col].map(lambda x: f"{float(x):.1%}")
        keep_cols = [col for col in ["LF", "Polarity", "Coverage", "Overlaps", "Conflicts"] if col in official_show.columns]
        add_table(
            doc,
            official_show,
            keep_cols,
            [{"LF": "LF", "Polarity": "极性", "Coverage": "覆盖率", "Overlaps": "重叠率", "Conflicts": "冲突率"}[c] for c in keep_cols],
            max_rows=12,
        )

    doc.add_heading("5. 文档差异与集中度", level=1)
    doc.add_paragraph(
        "不同文档长度、结构化程度和医学主题会影响 Lv1 的命中数量。下图按 accepted 信号数展示 Top 文档，文档名使用编号避免图中文字过密。"
    )
    add_picture_with_caption(doc, chart_paths["doc_concentration"], "图 4：accepted 信号数量最高的文档")
    top_show = top_docs.copy()
    top_show["labels_per_chunk"] = top_show["labels_per_chunk"].map(lambda x: f"{x:.2f}")
    add_table(
        doc,
        top_show,
        ["doc_code", "doc_name", "chunk_count", "accepted_label_rows", "labels_per_chunk"],
        ["编号", "文档名", "chunk数", "接受标签数", "每chunk标签数"],
        max_rows=12,
    )

    doc.add_heading("6. 质量观察", level=1)
    doc.add_paragraph(
        "质量观察不等同于错误结论，但它能帮助决定下一轮数据复核优先级。"
        "本批结果中，最需要关注的是空输出文档，以及 accepted 但缺少 supporting_lfs 或 evidence_texts 的记录。"
    )
    if math.isclose(metrics["accepted_rate"], 1.0):
        doc.add_paragraph(
            "另外，虽然 LF 热力图显示不同 LF 对不同标签的命中模式差异明显，但最终融合结果对所有有效 chunk 的所有标签均给出 accepted。"
            "这通常意味着融合阈值、majority fallback 或缺省概率需要复核，否则 Lv1 会失去“筛掉明显无关标签”的作用。"
        )
    add_picture_with_caption(doc, chart_paths["quality_flags"], "图 5：accepted 记录中的证据质量复核项")
    empty_docs = doc_summary[doc_summary["is_empty_output"] == 1][["doc_name", "chunk_file_bytes", "lf_file_bytes"]]
    if not empty_docs.empty:
        doc.add_heading("空输出文档清单", level=2)
        add_table(doc, empty_docs, ["doc_name", "chunk_file_bytes", "lf_file_bytes"], ["文档名", "chunk文件字节", "LF文件字节"], max_rows=30)
    add_bullet(doc, "空输出文档建议先检查 PDF 解析、chunk 生成和 Lv1 入口是否跳过。")
    add_bullet(doc, "accepted 但无支持 LF 的记录建议作为规则阈值或 majority fallback 的重点抽样对象。")
    add_bullet(doc, "对于高覆盖 LF，建议抽样核验证据文本是否能回链到原 chunk，避免高覆盖但证据松散。")

    doc.add_heading("7. 汇报建议", level=1)
    add_bullet(doc, "对外汇报时可强调：Lv1 已形成可追溯的 chunk-label 矩阵，能支撑后续实体抽取候选召回。")
    add_bullet(doc, "不要只汇报总体数量，应同时展示 LF 覆盖率和证据质量复核项，这样更能说明弱监督体系是否可靠。")
    add_bullet(doc, "下一步优先处理空输出文档、无证据 accepted 记录，以及标签间覆盖不均衡问题。")

    doc.save(out_docx)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", type=Path, required=True)
    parser.add_argument("--out-dir", type=Path, required=True)
    args = parser.parse_args()

    args.out_dir.mkdir(parents=True, exist_ok=True)
    charts_dir = args.out_dir / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    tables_dir = args.out_dir / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)

    configure_plot()
    official_available, official_reason = try_official_lfanalysis_available()
    chunk_df, lf_df, doc_df = load_dataset(args.input_dir)
    if chunk_df.empty or lf_df.empty:
        raise SystemExit("No usable Lv1 rows found.")

    label_summary = build_label_summary(chunk_df)
    lf_summary = build_lf_summary(lf_df)
    official_lf_summary = None
    if official_available:
        official_lf_summary = build_official_lfanalysis_summary(lf_df)
    doc_summary = build_doc_summary(chunk_df, doc_df)

    save_table(label_summary, tables_dir / "label_summary.csv")
    save_table(lf_summary, tables_dir / "lfanalysis_style_summary.csv")
    if official_lf_summary is not None:
        save_table(official_lf_summary, tables_dir / "official_snorkel_lfanalysis_summary.csv")
    save_table(doc_summary, tables_dir / "document_summary.csv")

    chart_paths = {
        "label_distribution": charts_dir / "label_distribution.png",
        "lf_coverage": charts_dir / "lf_coverage.png",
        "lf_heatmap": charts_dir / "lf_label_heatmap.png",
        "doc_concentration": charts_dir / "doc_concentration.png",
        "quality_flags": charts_dir / "quality_flags.png",
    }
    chart_label_distribution(label_summary, chart_paths["label_distribution"])
    chart_lf_coverage(lf_summary, chart_paths["lf_coverage"])
    chart_lf_label_heatmap(lf_df, chart_paths["lf_heatmap"])
    top_docs = chart_doc_concentration(doc_summary, chart_paths["doc_concentration"])
    chart_quality_flags(label_summary, chart_paths["quality_flags"])
    save_table(top_docs, tables_dir / "top_documents.csv")

    report_path = args.out_dir / "lv1_chunk标签结果与LFAnalysis分析报告.docx"
    build_report(
        out_docx=report_path,
        input_dir=args.input_dir,
        charts_dir=charts_dir,
        chunk_df=chunk_df,
        lf_df=lf_df,
        doc_summary=doc_summary,
        label_summary=label_summary,
        lf_summary=lf_summary,
        official_lf_summary=official_lf_summary,
        top_docs=top_docs,
        official_available=official_available,
        official_reason=official_reason,
        chart_paths=chart_paths,
    )

    summary = {
        "report_path": str(report_path),
        "official_lfanalysis_available": official_available,
        "official_lfanalysis_reason": official_reason,
        "documents": int(len(doc_summary)),
        "non_empty_documents": int((doc_summary["is_empty_output"] == 0).sum()),
        "empty_documents": int((doc_summary["is_empty_output"] == 1).sum()),
        "chunk_label_rows": int(len(chunk_df)),
        "lf_rows": int(len(lf_df)),
        "accepted_rows": int(((chunk_df["present"] == True) & (chunk_df["status"] == "accepted")).sum()),  # noqa: E712
    }
    (args.out_dir / "analysis_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
